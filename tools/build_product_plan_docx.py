from __future__ import annotations

import os
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "SkyCast_AI产品全链路规划书_v1.0.docx"

FONT_CN = "Microsoft YaHei"
FONT_EN = "Calibri"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "17324D"
INK = "1F2937"
MUTED = "64748B"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
GREEN = "176B4D"
GOLD = "7A5A00"
RED = "9B1C1C"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa: Sequence[int], indent_dxa: int = TABLE_INDENT_DXA) -> None:
    if sum(widths_dxa) != TABLE_WIDTH_DXA:
        raise ValueError(f"Column widths must sum to {TABLE_WIDTH_DXA}, got {sum(widths_dxa)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")


def set_run_font(run, size: float | None = None, bold: bool | None = None,
                 color: str | None = None, italic: bool | None = None,
                 font_name: str = FONT_CN) -> None:
    run.font.name = font_name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_EN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_EN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_style(style, size: float, color: str, bold: bool,
                    before: float, after: float, line_spacing: float = 1.1) -> None:
    style.font.name = FONT_CN
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_EN)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_EN)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_CN)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line_spacing
    style.paragraph_format.keep_with_next = True


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color=MUTED)


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_body(doc: Document, text: str, bold_prefix: str | None = None,
             italic: bool = False, color: str = INK):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_prefix and text.startswith(bold_prefix):
        first, rest = text[:len(bold_prefix)], text[len(bold_prefix):]
        r1 = p.add_run(first)
        set_run_font(r1, size=11, bold=True, color=color)
        r2 = p.add_run(rest)
        set_run_font(r2, size=11, color=color, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=color, italic=italic)
    return p


def add_bullet(doc: Document, text: str, level: int = 0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)
    return p


def create_decimal_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    level.append(lvl_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "120")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    abstract.append(level)
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(numbering.index(first_num), abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_numbered_group(doc: Document, items: Sequence[str]) -> None:
    num_id = create_decimal_numbering(doc)
    for text in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.167
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_pr.append(ilvl)
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.append(num_id_node)
        p._p.get_or_add_pPr().append(num_pr)
        r = p.add_run(text)
        set_run_font(r, size=11, color=INK)


def add_callout(doc: Document, label: str, text: str, fill: str = CALLOUT,
                accent: str = BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    r1 = p.add_run(f"{label}  ")
    set_run_font(r1, size=11, bold=True, color=accent)
    r2 = p.add_run(text)
    set_run_font(r2, size=11, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc: Document, headers: Sequence[str], rows: Iterable[Sequence[str]],
              widths_dxa: Sequence[int], header_fill: str = LIGHT_GRAY,
              center_cols: set[int] | None = None):
    rows = list(rows)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=9.5, bold=True, color=NAVY)
    for row_data in rows:
        row = table.add_row()
        for idx, text in enumerate(row_data):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            if center_cols and idx in center_cols:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(text))
            set_run_font(r, size=9.3, color=INK)
    set_table_geometry(table, widths_dxa)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(4)
    return table


def add_code_block(doc: Document, lines: Sequence[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    for idx, line in enumerate(lines):
        r = p.add_run(line)
        set_run_font(r, size=9, color=NAVY, font_name="Consolas")
        if idx < len(lines) - 1:
            r.add_break()
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_section_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build_document() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    configure_style(styles["Normal"], 11, INK, False, 0, 6, 1.10)
    configure_style(styles["Heading 1"], 16, BLUE, True, 16, 8, 1.0)
    configure_style(styles["Heading 2"], 13, BLUE, True, 12, 6, 1.0)
    configure_style(styles["Heading 3"], 12, DARK_BLUE, True, 8, 4, 1.0)
    for list_style in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[list_style]
        style.font.name = FONT_CN
        style.font.size = Pt(11)
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167
    styles["List Bullet"].paragraph_format.left_indent = Inches(0.5)
    styles["List Bullet"].paragraph_format.first_line_indent = Inches(-0.25)
    styles["List Bullet 2"].paragraph_format.left_indent = Inches(0.75)
    styles["List Bullet 2"].paragraph_format.first_line_indent = Inches(-0.25)
    styles["List Number"].paragraph_format.left_indent = Inches(0.5)
    styles["List Number"].paragraph_format.first_line_indent = Inches(-0.25)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run("SkyCast｜AI 产品全链路规划书")
    set_run_font(r, size=8.5, color=MUTED)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("SkyCast Product Blueprint  ·  第 ")
    set_run_font(r, size=8.5, color=MUTED)
    add_page_field(fp)
    r = fp.add_run(" 页")
    set_run_font(r, size=8.5, color=MUTED)

    # Cover - editorial_cover pattern.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(80)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PRODUCT STRATEGY & DELIVERY BLUEPRINT")
    set_run_font(r, size=10.5, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SkyCast（逐光）")
    set_run_font(r, size=30, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI 气象与天文景观决策助手")
    set_run_font(r, size=18, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(44)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("产品实现路径 · 技术路线 · AI 调用 · 用户验证 · 模型评估 · 研发治理")
    set_run_font(r, size=11, color=MUTED)

    add_callout(
        doc,
        "文档目的",
        "以成熟 AI 产品经理的方式，将概念收敛为可验证、可开发、可复盘的 0-1 产品方案；同时作为后续 Android/C++ 开发、GitHub 项目管理和秋招作品集的统一基线。",
        fill=LIGHT_BLUE,
        accent=DARK_BLUE,
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(50)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in ("版本：v1.0", "日期：2026 年 8 月 24 日", "状态：M0 产品与工程基线"):
        r = p.add_run(line)
        set_run_font(r, size=10, color=MUTED)
        r.add_break()

    add_section_break(doc)

    # Executive summary
    add_heading(doc, "执行摘要", 1)
    add_body(doc, "SkyCast 不应被定义为另一款通用天气 App，而应被定义为“面向户外影像场景的 AI 决策产品”。它需要回答五个问题：是否值得出发、何时拍、去哪里、将发生什么变化、是否安全。")
    add_callout(doc, "核心建议", "先用 4 周完成一个城市、两个历史事件、两条纵向切片；用 8 周形成可安装 APK、模型评估、C++ 性能报告和完整 GitHub 研发记录。不要以全国覆盖、社区、商业化或端到端大模型为首期目标。", LIGHT_BLUE, DARK_BLUE)
    add_heading(doc, "关键产品决策", 2)
    for item in [
        "目标用户首先是有明确拍摄决策需求的风光摄影师，而不是所有天气用户。",
        "首条纵向切片选择“城市晚霞机会卡”；第二条选择“雷达回放与 0-60 分钟短临研判”。",
        "LLM 只解释经过校验的结构化事实，不直接预测天气、不修改官方预警、不绕过安全策略。",
        "追风功能改为“安全观测点与撤离路线”，危险区和官方预警是硬约束。",
        "模型是否上线由离线基线、概率校准和真实用户反馈共同决定。",
        "Android 负责交互和离线体验；Python 负责数据接入与模型；C++ 负责可测量的高性能栅格、几何和渲染能力。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "阶段性成果", 2)
    add_table(
        doc,
        ["阶段", "要证明的问题", "核心交付物", "发布门槛"],
        [
            ("M0", "方案是否可执行", "PRD、架构、代码基线", "范围和验收清晰"),
            ("M1", "数据是否可获得", "事件包、数据契约", "可追溯、可回放"),
            ("M2", "用户链路是否成立", "晚霞纵向切片", "3 分钟完成决策"),
            ("M3", "模型是否有增益", "基线与评估报告", "优于简单基线"),
            ("M4", "作品是否可面试", "APK、视频、GitHub", "可复现、可讲述"),
        ],
        [900, 2300, 3000, 3160],
        center_cols={0},
    )

    add_heading(doc, "目录", 1)
    toc_items = [
        "1. 产品战略与边界", "2. 用户研究与需求验证", "3. 用户、场景与完整旅程",
        "4. 功能架构与版本范围", "5. 核心功能详细设计", "6. 数据产品与遥感融合",
        "7. AI 调用与智能决策架构", "8. 模型算法与评估体系", "9. 技术架构总览",
        "10. Android 与 C++ 客户端路线", "11. 后端、数据与工程化路线", "12. 合规、安全与风险治理",
        "13. 研发实施路线图", "14. 产品运营与指标体系", "15. 测试、验收与发布门槛",
        "16. 项目管理与 GitHub 实战", "17. 秋招作品集与面试叙事", "18. 下一步执行清单",
        "附录 A：核心数据对象", "附录 B：首批 API", "附录 C：参考资料与术语",
    ]
    for item in toc_items:
        add_bullet(doc, item)

    # 1 Strategy
    add_section_break(doc)
    add_heading(doc, "1. 产品战略与边界", 1)
    add_heading(doc, "1.1 产品愿景", 2)
    add_body(doc, "让用户在复杂、多源、具有不确定性的气象信息中，快速得到可解释、可追溯且安全的户外影像决策。")
    add_heading(doc, "1.2 一句话定位", 2)
    add_callout(doc, "定位", "SkyCast 是面向风光、天文和强天气摄影用户的 AI 气象决策助手，融合雷达、卫星、地面气象、天文、地形与道路数据，给出拍摄概率、最佳时间、候选地点和安全建议。")
    add_heading(doc, "1.3 用户价值主张", 2)
    add_table(
        doc,
        ["用户问题", "传统产品答案", "SkyCast 答案"],
        [
            ("今晚会不会有晚霞？", "云量、天气描述", "概率、染色窗口、方向、原因"),
            ("去哪里拍？", "地图和通用 POI", "方向无遮挡、景观潜力、可达性排序"),
            ("雷暴往哪里走？", "雷达回放", "15-60 分钟概率范围与不确定性"),
            ("是否安全？", "官方预警", "官方预警 + AI 风险研判 + 硬约束"),
            ("预测为什么错？", "通常无法解释", "时间轴回放、数据版本和结果反馈"),
        ],
        [2300, 2800, 4260],
    )
    add_heading(doc, "1.4 产品原则", 2)
    for item in [
        "决策优先：展示信息必须服务于具体行动。",
        "可解释优先：分数必须拆解为影响因素和不确定性。",
        "安全优先：风险过滤优先级高于景观分数。",
        "基线优先：先建立可解释基线，再尝试复杂模型。",
        "回放优先：实时链路必须能够用历史事件复现。",
        "合规优先：数据来源、模型身份和官方预警必须清晰分层。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "1.5 商业与作品集边界", 2)
    add_body(doc, "秋招阶段的目标不是证明商业规模，而是证明产品判断、技术完整性和工程执行力。全国覆盖、付费订阅、摄影社区和商业导航属于商业化阶段，不应挤占 MVP 的验证资源。")

    # 2 User research
    add_heading(doc, "2. 用户研究与需求验证", 1)
    add_heading(doc, "2.1 研究目标", 2)
    add_body(doc, "研究不是询问用户“喜不喜欢这个功能”，而是验证其现有决策行为、成本、失败原因和付费/迁移意愿。")
    add_heading(doc, "2.2 访谈样本", 2)
    add_table(
        doc,
        ["用户组", "建议人数", "筛选条件", "重点验证"],
        [
            ("城市风光摄影师", "4-6", "每月拍摄至少 2 次", "晚霞判断、地点选择"),
            ("天文摄影用户", "3-4", "有银河或月升拍摄经验", "云量、光污染、地平线"),
            ("强天气摄影用户", "2-3", "看过雷达且有外拍经验", "短临信息、安全边界"),
            ("普通天气用户", "2-3", "偶尔拍摄", "理解成本、是否过度专业"),
        ],
        [2100, 1100, 3000, 3160],
        center_cols={1},
    )
    add_heading(doc, "2.3 访谈问题", 2)
    add_numbered_group(doc, [
        "最近一次为了朝霞、晚霞或天文景观出门是什么时候？请按时间复盘全过程。",
        "出发前看了哪些 App、群、雷达或指标？哪一步最耗时？",
        "最近一次判断错误是什么原因？损失了多少时间或交通成本？",
        "你如何选择拍摄地点？会因为哪些风险取消？",
        "如果只能保留一个功能，你最希望它替你完成哪个判断？",
        "什么证据会让你相信一个 AI 分数？什么情况会让你卸载？",
    ])
    add_heading(doc, "2.4 假设与验证方法", 2)
    add_table(
        doc,
        ["产品假设", "验证方式", "通过标准", "失败后的调整"],
        [
            ("用户愿意为明确的拍摄窗口打开 App", "原型任务测试", "8 人中至少 6 人完成决策", "简化首页与术语"),
            ("分解原因比单一分数更可信", "双版本访谈", "解释版信任评分更高", "改为证据卡而非长文"),
            ("历史回放能建立信任", "盲测 10 个事件", "用户能理解成败原因", "强化时间轴和基线"),
            ("强天气用户接受安全硬约束", "情景测试", "无用户选择穿越危险区", "取消路线功能或扩大缓冲"),
        ],
        [2300, 1900, 2300, 2860],
    )
    add_callout(doc, "研究纪律", "不使用“你会不会用”“你愿不愿意付费”作为主要证据；优先观察真实行为、历史记录、现有工具链和已经发生的成本。", fill="FFF8E8", accent=GOLD)

    # 3 Personas & journeys
    add_heading(doc, "3. 用户、场景与完整旅程", 1)
    add_heading(doc, "3.1 核心 Persona", 2)
    add_table(
        doc,
        ["Persona", "目标", "现有行为", "核心痛点", "首要价值"],
        [
            ("城市逐光者", "下班后拍晚霞", "看天气、云图、群消息", "判断分散、容易错过", "15 秒完成机会判断"),
            ("计划型天文用户", "提前规划银河/月升", "查星图、天气、光污染", "工具多且地点验证难", "统一时间地点决策"),
            ("专业追风摄影者", "拍摄雷暴结构", "雷达、预警、地图并用", "路径不确定且风险高", "短临概率与安全边界"),
        ],
        [1500, 1800, 2000, 2100, 1960],
    )
    add_heading(doc, "3.2 核心 JTBD", 2)
    add_callout(doc, "Job to be Done", "当我考虑为一次景观拍摄投入时间和交通成本时，我希望快速判断成功概率、最佳窗口、地点与风险，从而决定出发、等待、换点或取消。")
    add_heading(doc, "3.3 用户全旅程", 2)
    stages = [
        ("发现", "首页机会卡", "是否值得进一步查看"),
        ("理解", "评分、概率、理由", "是否相信结论"),
        ("计划", "地点、方向、到达时间", "选择候选点"),
        ("执行", "雷达变化、提醒", "等待、换点或取消"),
        ("安全", "预警、风险和撤离", "停止拍摄并退出"),
        ("复盘", "事件回放、结果反馈", "理解模型成败"),
    ]
    add_table(doc, ["阶段", "关键触点", "用户决策"], stages, [1200, 3600, 4560], center_cols={0})

    # 4 Feature architecture
    add_heading(doc, "4. 功能架构与版本范围", 1)
    add_heading(doc, "4.1 功能地图", 2)
    add_table(
        doc,
        ["模块", "P0：面试 MVP", "P1：完整作品版", "P2：商业探索"],
        [
            ("天气首页", "城市、实况、逐小时、机会卡", "个性化订阅", "多端与会员"),
            ("朝晚霞", "评分、窗口、方向、解释", "城市热力图、地点排序", "个性化校准"),
            ("雷达", "历史动画、时间轴", "实时接入、C++渲染", "全国高可用"),
            ("强对流", "光流基线、概率锥", "单体追踪、集合预报", "专业模式"),
            ("天文", "基础日月信息", "银河/月升/星轨评估", "器材与构图建议"),
            ("地点路线", "候选地点样例", "安全过滤、导航接入", "合作 POI"),
            ("反馈", "拍到/未拍/取消", "原因标签、模型校准", "社区内容"),
        ],
        [1500, 2800, 2800, 2260],
    )
    add_heading(doc, "4.2 MVP 取舍", 2)
    add_body(doc, "首版只做一个城市和固定历史事件包，以保证数据、模型、客户端和反馈形成完整闭环。用户看不到全国覆盖，但面试官能看到完整产品方法。")
    for item in [
        "必须有：一条晚霞端到端链路、一条雷达回放与短临链路。",
        "应该有：离线缓存、来源时间、置信度、模型版本和异常降级。",
        "可以有：候选地点样例和地图导航入口。",
        "明确不做：社区、支付、自研导航、持续后台定位和模型自动发布预警。",
    ]:
        add_bullet(doc, item)

    # 5 Detailed functional design
    add_heading(doc, "5. 核心功能详细设计", 1)
    add_heading(doc, "5.1 城市首页", 2)
    add_body(doc, "首页不是天气信息堆叠页，而是决策入口。首屏只保留城市、当前状态、今日场景机会和风险提醒。")
    add_table(
        doc,
        ["元素", "内容", "用户动作", "验收标准"],
        [
            ("场景机会卡", "晚霞 82 分 / 76%", "查看详情", "15 秒理解是否值得拍"),
            ("时间窗口", "18:43-19:07", "加入提醒", "含峰值和误差范围"),
            ("风险卡", "官方预警/数据过期", "查看来源", "优先级高于机会卡"),
            ("数据状态", "更新时间、来源", "刷新", "不把缓存当实时"),
        ],
        [1800, 2300, 1800, 3460],
    )
    add_heading(doc, "5.2 朝霞晚霞评估", 2)
    add_body(doc, "输出必须包含评分、概率、时间、方向、置信度和可解释因素。首版不承诺准确预测天空颜色，而是评估“出现可拍摄染色景观的概率”。")
    add_code_block(doc, [
        "晚霞评分：82 / 100          成功概率：76%",
        "预计染色：18:43 - 19:07     峰值：18:51 - 18:59",
        "最佳方向：西偏北 12°       置信度：中等",
        "有利：西侧地平线通透、中高云适量",
        "限制：低层湿度偏高、云系移动仍有不确定性",
        "数据时间：18:20             模型：sunset-rules-v0",
    ])
    add_heading(doc, "5.3 场景地图与地点排序", 2)
    for item in [
        "地图展示城市场景热力格点，而不是把每个格点解释成精确预报。",
        "候选地点使用视野方向、DEM 地平线、景观元素、可达性、安全性和到达时间综合排序。",
        "内部使用 WGS84；进入高德显示和导航边界时转换为 GCJ-02。",
        "地点详情展示推荐原因、限制条件、数据时间和停车/步行信息。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "5.4 雷达回放与气象追溯", 2)
    add_body(doc, "事件时间轴统一承载雷达帧、卫星云图、地面实况、官方预警和模型输出，使用户能够回看某个时刻系统当时掌握了什么，而不是事后使用未来数据解释过去。")
    add_heading(doc, "5.5 强对流短临研判", 2)
    add_body(doc, "产品输出 15/30/45/60 分钟概率范围和风暴单体移动趋势。不得把概率锥呈现为确定路线，也不得把 AI 结果命名为气象灾害预警。")
    add_heading(doc, "5.6 安全观测点与撤离路线", 2)
    add_body(doc, "路线目标函数可以优化景观价值、到达时间和可达性，但官方预警区、预测核心缓冲区、山洪风险区、封闭区域和夜间不可达区域是不可放松的硬约束。")
    add_callout(doc, "安全红线", "任何时候只要安全策略与景观推荐冲突，系统必须取消推荐并解释原因。产品不提供“追入风暴核心”的路线。", fill="FDECEC", accent=RED)
    add_heading(doc, "5.7 天文景观评估", 2)
    add_body(doc, "P1 阶段增加银河、月升月落、星轨和流星雨场景。输入包括太阳/月亮高度角、月相、天文暮光、云量、透明度、光污染、地平线遮挡和目标方位。输出仍采用概率、窗口、方向和解释结构。")
    add_heading(doc, "5.8 结果反馈", 2)
    for item in [
        "一级反馈：拍到了、没拍到、没有出发。",
        "二级原因：云太厚、地平线被挡、时间偏差、天气突变、地点不可达、主动取消。",
        "可选证据：用户照片、拍摄时间、方向；默认不公开。",
        "反馈进入训练前必须做去重、时空校验和隐私处理。",
    ]:
        add_bullet(doc, item)

    # 6 Data
    add_heading(doc, "6. 数据产品与遥感融合", 1)
    add_heading(doc, "6.1 数据分工", 2)
    add_table(
        doc,
        ["数据类型", "主要用途", "更新/特性", "首期策略"],
        [
            ("天气预报与实况", "温湿风、云量、能见度", "按供应商", "授权 API + 适配器"),
            ("天气雷达", "回波、单体、短临", "组网产品约 6 分钟", "历史事件包先行"),
            ("FY-4 卫星", "云类型、云顶、云系移动", "部分产品约 15 分钟", "P1 实时融合"),
            ("DEM", "真实地平线、地形风险", "静态", "离线预处理"),
            ("Sentinel/土地覆盖", "水体、植被、景观潜力", "非实时", "地点静态评分"),
            ("官方预警", "风险与通知", "事件驱动", "原文传播、不可改写"),
            ("地图道路", "POI、路线、可达性", "服务更新", "高德 SDK/API"),
        ],
        [1700, 2800, 2100, 2760],
    )
    add_heading(doc, "6.2 数据准入表", 2)
    add_body(doc, "任何数据接入前必须登记：所有者、授权范围、商业展示权、更新频率、平均/尾部延迟、坐标系、时间语义、缺测规则、存储期限、费用和替代方案。没有明确授权的数据只能用于本地研究样例。")
    add_heading(doc, "6.3 标准化处理", 2)
    for item in [
        "时间：UTC 存储，明确 observedAt、issuedAt、ingestedAt、expiresAt。",
        "空间：WGS84 作为内部标准，并保留原始 CRS 和转换记录。",
        "质量：处理缺帧、重复、乱序、坏值、遮挡和投影偏移。",
        "血缘：每个模型结果保留输入对象 ID、校验和、模型版本和请求 ID。",
        "回放：历史事件使用与实时相同的数据对象和服务接口。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "6.4 数据新鲜度策略", 2)
    add_table(
        doc,
        ["数据", "Fresh", "Degraded", "Expired 后动作"],
        [
            ("雷达", "≤6 分钟", "7-12 分钟", "停止短临和路线建议"),
            ("卫星", "≤15 分钟", "16-30 分钟", "降级为预报云量"),
            ("短时预报", "≤60 分钟", "61-180 分钟", "只显示缓存且标记"),
            ("官方预警", "有效期内", "临近过期", "移入历史记录"),
        ],
        [1800, 1800, 2200, 3560],
        center_cols={1, 2},
    )
    add_body(doc, "上述阈值是首期产品策略，不是永久气象标准；上线前应根据真实数据延迟分布和业务风险校准。")

    # 7 AI
    add_heading(doc, "7. AI 调用与智能决策架构", 1)
    add_heading(doc, "7.1 AI 能力分层", 2)
    add_table(
        doc,
        ["层", "职责", "建议技术", "禁止事项"],
        [
            ("数值计算层", "太阳几何、地形、统计特征", "C++/Python确定性算法", "由 LLM 代算关键数值"),
            ("预测模型层", "评分、概率、光流、排序", "LightGBM/pySTEPS/PyTorch", "无基线上线复杂模型"),
            ("策略层", "安全过滤、过期降级", "规则引擎", "被模型自由绕过"),
            ("LLM解释层", "摘要、问答、原因组织", "结构化提示 + JSON Schema", "生成官方预警或新概率"),
        ],
        [1600, 2300, 2500, 2960],
    )
    add_heading(doc, "7.2 LLM 调用链", 2)
    add_code_block(doc, [
        "用户请求",
        "  → 权限与场景识别",
        "  → 聚合结构化气象事实",
        "  → 数据新鲜度与安全策略检查",
        "  → 生成模型输出（score/probability/reasons）",
        "  → LLM 只对上述事实进行解释",
        "  → JSON Schema 校验 + 敏感词/安全校验",
        "  → 模板化降级或返回客户端",
        "  → 记录 promptVersion/modelVersion/requestId",
    ])
    add_heading(doc, "7.3 LLM 输入契约", 2)
    add_body(doc, "LLM 不读取未经预处理的 2048×2048 雷达图，也不从模糊截图自行判断风暴。输入必须是经过数值模型和策略层校验的结构化事实。")
    add_code_block(doc, [
        "{",
        "  sceneType: 'sunset',",
        "  location: { cityId, lat, lon },",
        "  validWindow: { start, peak, end },",
        "  assessment: { score, probability, confidence },",
        "  favorableFactors: [...],",
        "  limitingFactors: [...],",
        "  dataStatus: [{ source, observedAt, freshness }],",
        "  officialAlerts: [...],",
        "  safetyDecision: { allowed, reasonCodes },",
        "  modelVersion: 'sunset-rules-v0'",
        "}",
    ])
    add_heading(doc, "7.4 LLM 输出契约", 2)
    add_body(doc, "输出固定为摘要、建议动作、不确定性和引用来源四部分。客户端不展示模型的隐藏推理过程，只展示可核验的简短原因。")
    add_code_block(doc, [
        "{",
        "  summary: string,",
        "  recommendedAction: 'go' | 'wait' | 'cancel' | 'inspect',",
        "  evidence: [{ factorCode, explanation }],",
        "  uncertainty: string,",
        "  sourceRefs: [string],",
        "  aiGenerated: true",
        "}",
    ])
    add_heading(doc, "7.5 模型路由、成本和降级", 2)
    for item in [
        "短摘要和批量地点解释使用成本较低的小模型；复杂问答才升级模型。",
        "按 sceneId + modelVersion + dataVersion 缓存解释，避免重复调用。",
        "设置输入/输出 token 上限、超时、重试和每日预算。",
        "结构化校验失败时不继续追问模型，直接使用确定性模板。",
        "模型服务不可用时不影响数值评分、官方预警和安全策略。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "7.6 AI Eval", 2)
    add_table(
        doc,
        ["评估维度", "测试集", "指标", "发布门槛"],
        [
            ("事实一致性", "结构化事实与回答对照", "字段一致率", "关键字段 100%"),
            ("来源完整性", "所有解释样例", "sourceRefs 覆盖率", "100%"),
            ("安全遵循", "危险情景红队集", "违规建议数", "0"),
            ("表达可懂", "目标用户任务测试", "理解正确率", "≥80%"),
            ("稳定性", "同输入重复调用", "动作一致率", "≥95%"),
            ("成本", "一周回放流量", "单决策成本", "满足预算上限"),
        ],
        [1800, 2600, 2300, 2660],
        center_cols={3},
    )

    # 8 Algorithms and validation
    add_heading(doc, "8. 模型算法与评估体系", 1)
    add_heading(doc, "8.1 朝晚霞模型路线", 2)
    add_body(doc, "第一阶段建立可解释规则分；第二阶段用历史标签训练 LightGBM 或逻辑回归并做概率校准；第三阶段再评估时空深度模型是否带来稳定增益。")
    for item in [
        "太阳几何：高度角、方位角、暮光阶段。",
        "云场：低云地平线遮挡、中高云覆盖、云顶和云系移动。",
        "大气：湿度、能见度、气溶胶、降水和风。",
        "地形：目标方向地平线、海拔、视野遮挡。",
        "标签：是否出现可拍染色、起峰止时间、用户结果反馈。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "8.2 强对流模型路线", 2)
    add_code_block(doc, [
        "雷达质控 → 阈值分割 → 连通域/轮廓 → 单体匹配",
        "         → 光流运动场 → 半拉格朗日外推 → 概率集合",
        "         → persistence baseline 对比 → 安全策略 → 客户端",
    ])
    add_body(doc, "地面风向不能直接当作风暴移动方向。首版以雷达回波运动为主；有可靠高空资料时再把引导气流作为辅助特征。")
    add_heading(doc, "8.3 评估数据切分", 2)
    for item in [
        "按天气过程和日期切分训练、验证和测试集，禁止相邻帧跨集合造成泄漏。",
        "按城市、季节和天气类型分别报告指标，避免总体平均掩盖失败场景。",
        "任何参数调优不得查看最终测试集。",
        "历史回放只使用当时可获得的数据，禁止未来信息穿越。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "8.4 指标与发布门槛", 2)
    add_table(
        doc,
        ["模型", "核心指标", "基线", "首期发布标准"],
        [
            ("晚霞概率", "Brier、AUC、Top-K", "城市气候概率", "Brier 至少改善 10%"),
            ("染色时间", "MAE", "日落固定偏移", "MAE ≤15 分钟（目标）"),
            ("雷达外推", "FSS/CSI/POD/FAR", "persistence", "30 分钟 FSS 改善 ≥10%"),
            ("单体轨迹", "中心位置误差", "匀速外推", "分 15/30/60 分钟报告"),
            ("概率锥", "可靠性、覆盖率", "固定缓冲区", "无系统性过度自信"),
        ],
        [1800, 2300, 2200, 3060],
    )
    add_body(doc, "若真实数据未达到目标，不伪造指标。作品版应展示失败分析、限制条件和下一步，而不是只展示最好案例。")

    # 9 Architecture
    add_heading(doc, "9. 技术架构总览", 1)
    add_code_block(doc, [
        "天气 / 雷达 / FY-4 / DEM / 地图 / 官方预警",
        "                    ↓",
        "Provider Adapters → 质控 → 标准观测对象",
        "                    ↓",
        "PostGIS + 对象存储 + Redis",
        "                    ↓",
        "场景评分 / 雷达短临 / 地点排序 / 安全策略",
        "                    ↓",
        "REST API + WebSocket + Push",
        "                    ↓",
        "Android Compose + Room + C++ meteocore",
    ])
    add_heading(doc, "9.1 架构原则", 2)
    for item in [
        "模块化单体起步，不在 MVP 引入微服务。",
        "Provider Adapter 隔离数据供应商，便于替换授权源。",
        "Room 是客户端 UI 的单一事实来源，网络成功后先落库再展示。",
        "栅格使用瓦片或对象存储，禁止把大矩阵塞入 JSON。",
        "算法、策略和 LLM 三层解耦，任何一层失败都能确定性降级。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "9.2 仓库建议", 2)
    add_code_block(doc, [
        "SkyCast/",
        "├─ docs/                  PRD、架构、模型卡、决策记录",
        "├─ core/                  C++20 共享核心",
        "├─ android/               Kotlin/Compose Android 客户端",
        "├─ server/                FastAPI 数据与模型服务",
        "├─ data/sample/           可公开历史事件样例",
        "├─ evaluation/            离线评估与报告",
        "└─ .github/               Issue、PR、CI、Release",
    ])

    # 10 Android/C++
    add_heading(doc, "10. Android 与 C++ 客户端路线", 1)
    add_heading(doc, "10.1 Android 技术栈", 2)
    add_table(
        doc,
        ["领域", "技术", "用途"],
        [
            ("UI", "Kotlin + Jetpack Compose", "单 Activity、响应式界面"),
            ("状态", "ViewModel + StateFlow", "单向数据流"),
            ("本地", "Room + DataStore", "离线读、用户偏好"),
            ("同步", "Retrofit/OkHttp + WorkManager", "缓存、重试、约束同步"),
            ("地图", "高德地图/导航 SDK", "地图、POI、路线"),
            ("Native", "NDK + CMake + JNI", "高性能计算与渲染"),
            ("测试", "JUnit + Compose UI Test + Benchmark", "功能与性能验证"),
        ],
        [1700, 3200, 4460],
    )
    add_heading(doc, "10.2 C++ 模块的合理边界", 2)
    add_body(doc, "C++ 必须用于性能敏感且可量化的路径，而不是为了简历而重写普通业务代码。建议共享模块命名为 meteocore。")
    for item in [
        "雷达栅格解码、颜色映射和等值线提取。",
        "雷达动画 LRU 帧缓存、预取和内存复用。",
        "OpenGL ES 自定义雷达图层渲染。",
        "太阳方位、DEM 地平线和空间几何计算。",
        "桌面与 Android 共用的数据新鲜度和领域校验。",
    ]:
        add_bullet(doc, item)
    jni_heading = add_heading(doc, "10.3 JNI 设计规则", 2)
    jni_heading.paragraph_format.space_after = Pt(3)
    for item in [
        "Kotlin 持有显式 native handle，并保证 create/destroy 成对。",
        "使用 DirectByteBuffer 或扁平数组批量传输，禁止逐像素 JNI。",
        "Native 异常在边界内转成状态码或受控异常。",
        "C++ 对象不保存失效的 JNI 局部引用。",
        "每个 Native API 都有桌面测试、Android 集成测试和 Benchmark。",
    ]:
        jni_bullet = add_bullet(doc, item)
        jni_bullet.paragraph_format.space_after = Pt(2)
        jni_bullet.paragraph_format.line_spacing = 1.0
    performance_heading = add_heading(doc, "10.4 客户端性能指标", 2)
    performance_heading.paragraph_format.space_before = Pt(6)
    performance_heading.paragraph_format.space_after = Pt(3)
    performance_table = add_table(
        doc,
        ["指标", "测量方法", "目标"],
        [
            ("雷达动画帧耗时", "Macrobenchmark p50/p95", "目标设备稳定 30fps"),
            ("峰值内存", "Profiler + 固定事件包", "无持续增长、无 OOM"),
            ("JNI 传输", "调用数与拷贝字节", "批量传输、可解释"),
            ("离线首屏", "断网启动", "展示最近成功同步结果"),
            ("崩溃", "测试设备回归", "主路径 0 crash"),
        ],
        [2200, 3600, 3560],
    )
    for cell in performance_table.rows[0].cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.keep_with_next = True

    # 11 backend
    add_heading(doc, "11. 后端、数据与工程化路线", 1)
    add_heading(doc, "11.1 后端模块", 2)
    add_code_block(doc, [
        "server/app/providers     外部数据适配器",
        "server/app/ingestion     调度、质控、标准化",
        "server/app/domain        领域模型和策略",
        "server/app/models        评分、光流、评估",
        "server/app/api           REST / WebSocket",
        "server/tests             单元、契约、回放测试",
    ])
    add_heading(doc, "11.2 服务与存储", 2)
    add_table(
        doc,
        ["组件", "首期选择", "职责"],
        [
            ("API", "FastAPI", "快速迭代、Schema、模型集成"),
            ("空间数据", "PostgreSQL + PostGIS", "POI、多边形、轨迹、查询"),
            ("栅格", "本地/对象存储", "雷达、卫星、瓦片"),
            ("缓存", "Redis（P1）", "热点结果和任务状态"),
            ("任务", "进程内调度→任务队列", "数据接入与模型运行"),
            ("部署", "Docker Compose", "本地复现和演示"),
        ],
        [1900, 2700, 4760],
    )
    add_heading(doc, "11.3 可观测性", 2)
    for item in [
        "结构化日志包含 requestId、dataVersion、modelVersion、耗时和降级原因。",
        "监控数据接入成功率、延迟分位数、缺帧率、API p95 和模型失败率。",
        "对每个事件保存可回放清单，保证线上问题可在本地重现。",
        "密钥通过环境变量或本地安全配置注入，禁止进入 Git。",
    ]:
        add_bullet(doc, item)

    # 12 compliance
    add_heading(doc, "12. 合规、安全与风险治理", 1)
    add_heading(doc, "12.1 官方预警与 AI 研判分层", 2)
    add_table(
        doc,
        ["内容", "身份", "允许操作", "展示要求"],
        [
            ("官方预警", "气象主管机构发布", "原文传播、跳转来源", "来源、时间、有效期、级别"),
            ("场景概率", "SkyCast AI 评估", "解释拍摄机会", "AI 标识、置信度、模型版本"),
            ("强对流风险", "实验性研判", "展示概率和范围", "不得命名为官方预警"),
            ("安全策略", "产品规则", "取消地点或路线", "说明触发规则"),
        ],
        [1800, 2300, 2600, 2660],
    )
    add_heading(doc, "12.2 隐私与权限", 2)
    for item in [
        "MVP 只在前台、用户主动操作时获取定位；不申请永久后台位置。",
        "通知权限在用户理解订阅价值后再请求。",
        "照片反馈默认私有，删除 EXIF 中不必要的精确位置和设备信息。",
        "账号、定位、反馈和日志分别设置最小保留期限。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "12.3 风险矩阵", 2)
    add_table(
        doc,
        ["风险", "概率", "影响", "应对", "Owner"],
        [
            ("雷达数据无商用授权", "高", "高", "Provider 抽象 + 历史样例 + 授权核实", "产品/数据"),
            ("模型看起来准但不校准", "中", "高", "概率可靠性与基线评估", "算法"),
            ("路线诱导进入危险区", "低", "极高", "硬约束、缓冲区、取消机制", "安全"),
            ("Android 后台能力受限", "中", "中", "前台行程、服务端推送", "客户端"),
            ("功能范围失控", "高", "中", "阶段门与非范围清单", "PM"),
            ("LLM 幻觉或改写事实", "中", "高", "Schema、引用、模板降级", "AI"),
        ],
        [2100, 900, 900, 3900, 1560],
        center_cols={1, 2, 4},
    )

    # 13 roadmap
    add_heading(doc, "13. 研发实施路线图", 1)
    add_heading(doc, "13.1 8 周实施计划", 2)
    add_table(
        doc,
        ["周", "产品目标", "代码实现", "验收证据"],
        [
            ("W1", "数据可行性", "领域对象、事件包读取", "两个事件可重放"),
            ("W2", "首条 API", "FastAPI 固定 Provider", "契约测试通过"),
            ("W3", "Android 纵向切片", "Compose + Room + API", "离线机会卡"),
            ("W4", "面试 MVP", "晚霞规则分 + 时间轴", "3 分钟演示"),
            ("W5", "短临基线", "persistence + 光流", "离线评估报告"),
            ("W6", "C++能力", "解码、缓存、JNI", "性能对比"),
            ("W7", "安全与地点", "策略过滤 + 地图路线", "危险情景测试"),
            ("W8", "发布", "CI、APK、Release", "v0.1.0 作品包"),
        ],
        [700, 2200, 3300, 3160],
        center_cols={0},
    )
    add_heading(doc, "13.2 每个功能的执行模板", 2)
    add_numbered_group(doc, [
        "PM 定义：用户问题、价值、范围、非范围和验收标准。",
        "数据确认：来源、时间语义、授权、缺测和替代方案。",
        "算法确认：基线、指标、切分、失败条件和降级。",
        "开发实现：Issue → 分支 → 小步提交 → 测试 → Pull Request。",
        "产品验收：按用户路径和异常路径验收，不以“代码写完”为完成。",
        "复盘发布：记录指标、偏差、决策和下一步。",
    ])
    add_heading(doc, "13.3 第一条纵向切片拆解", 2)
    add_table(
        doc,
        ["任务", "输入", "输出", "完成标准"],
        [
            ("历史事件包", "人工整理数据", "JSON + 栅格清单", "可重复读取"),
            ("SceneAssessment", "结构化特征", "评分契约", "非法值被拒绝"),
            ("固定 Provider", "eventId", "API 响应", "契约测试通过"),
            ("Room 缓存", "API 响应", "本地实体", "断网可显示"),
            ("机会卡", "本地 StateFlow", "Compose UI", "来源和时间可见"),
            ("结果反馈", "用户选择", "Feedback 记录", "可导出评估"),
        ],
        [2000, 2300, 2200, 2860],
    )

    # 14 metrics
    add_heading(doc, "14. 产品运营与指标体系", 1)
    add_heading(doc, "14.1 北极星指标", 2)
    add_callout(doc, "有效拍摄决策数", "用户查看评估后执行出发、等待、换点或取消中的一种决策，并在事件后完成结果反馈的次数。该指标同时要求产品产生行动价值和反馈闭环。")
    add_heading(doc, "14.2 指标树", 2)
    add_table(
        doc,
        ["层级", "指标", "解释"],
        [
            ("触达", "机会卡曝光、通知到达", "是否进入决策入口"),
            ("理解", "详情查看率、解释展开率", "是否需要更多证据"),
            ("行动", "地点查看、路线发起、提醒订阅", "是否形成真实意图"),
            ("结果", "拍到/未拍/取消", "产品建议的结果"),
            ("闭环", "反馈完成率、7 日留存", "是否积累长期价值"),
            ("质量", "Brier/FSS/数据新鲜度", "模型与工程质量"),
            ("安全", "违规建议、危险区穿越", "必须独立监控且为 0"),
        ],
        [1500, 3400, 4460],
    )
    add_heading(doc, "14.3 埋点原则", 2)
    for item in [
        "事件名称描述用户行为，不描述 UI 组件实现。",
        "所有评估相关事件携带 sceneId、modelVersion 和 dataFreshness。",
        "不采集与产品决策无关的高精度位置轨迹。",
        "指标必须区分历史回放和实时使用，避免演示流量污染。",
    ]:
        add_bullet(doc, item)

    # 15 QA
    add_heading(doc, "15. 测试、验收与发布门槛", 1)
    add_heading(doc, "15.1 测试金字塔", 2)
    add_table(
        doc,
        ["层", "测试内容", "工具/方式"],
        [
            ("单元", "领域校验、时间、坐标、策略", "C++/Kotlin/Python 单元测试"),
            ("契约", "Provider、API、Schema", "固定样例与快照"),
            ("回放", "完整历史事件链", "事件包自动回归"),
            ("UI", "首页、时间轴、离线和错误态", "Compose UI Test"),
            ("性能", "帧耗时、内存、JNI、API", "Benchmark/Profiler"),
            ("AI Eval", "事实、安全、稳定、成本", "金标集 + 红队集"),
            ("用户", "3 分钟任务与理解", "可用性测试"),
        ],
        [1500, 4400, 3460],
    )
    add_heading(doc, "15.2 Release Gate", 2)
    for item in [
        "主路径自动测试通过，历史事件可重放。",
        "所有结论显示来源、时间、置信度和模型版本。",
        "过期、缺帧、乱序和服务失败均有明确降级。",
        "模型达到预设基线；否则只进入实验入口。",
        "危险情景测试中违规路线和误导建议为 0。",
        "APK、README、演示脚本和性能/模型报告齐全。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "15.3 用户验收脚本", 2)
    add_numbered_group(doc, [
        "给用户一个城市和日期，请其判断是否出发，并口述依据。",
        "让用户从三个地点中选择一个并解释原因。",
        "展示雷达数据过期状态，观察用户是否仍误以为是实时。",
        "展示官方预警与高景观分冲突，检查用户是否理解取消原因。",
        "完成事件后提交反馈，记录流程时间和困惑点。",
    ])

    # 16 Project management
    add_heading(doc, "16. 项目管理与 GitHub 实战", 1)
    add_heading(doc, "16.1 双角色协作", 2)
    add_body(doc, "你同时承担 Product Owner 和学习者角色；Codex 承担产品经理、架构师和实现协作者角色。所有真实产品代码由 Codex 分步实现并解释；Git、GitHub、环境配置和产品操作由你亲手执行。")
    add_heading(doc, "16.2 GitHub Flow", 2)
    add_code_block(doc, [
        "Product Epic → GitHub Issue → feat/编号-名称 分支",
        "             → 小步 Conventional Commit",
        "             → Pull Request + 测试证据",
        "             → 产品验收 → 合并 main → Release Tag",
    ])
    add_heading(doc, "16.3 Definition of Done", 2)
    for item in [
        "用户故事、范围和验收标准已更新。",
        "实现通过自动测试、静态检查和异常路径验证。",
        "数据来源、模型版本和降级状态可见。",
        "README/架构/模型卡按影响更新。",
        "PR 含截图、测试输出或性能对比。",
        "合并后 main 可构建，Issue 被关闭。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "16.4 研发节奏", 2)
    add_table(
        doc,
        ["节奏", "动作", "产物"],
        [
            ("每日", "选择一个可验收小任务", "代码 + 测试 + 提交"),
            ("每 2-3 日", "完成一条小 PR", "可运行增量"),
            ("每周", "Demo 与指标复盘", "周报、风险、下周范围"),
            ("每里程碑", "阶段门验收", "Tag、视频、决策记录"),
        ],
        [1600, 3900, 3860],
    )

    # 17 Portfolio
    add_heading(doc, "17. 秋招作品集与面试叙事", 1)
    add_heading(doc, "17.1 AI 产品经理叙事", 2)
    for item in [
        "从通用天气功能收敛到摄影决策场景的过程。",
        "如何用用户研究、假设和阶段门控制范围。",
        "为什么数值模型负责预测、LLM 只负责解释。",
        "如何定义离线指标、在线指标、概率校准和安全评估。",
        "如何处理数据授权、官方预警和高风险路线。",
        "如何从失败案例更新 PRD 与模型。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "17.2 C++ Android 客户端叙事", 2)
    for item in [
        "C++ 模块为何只覆盖栅格、空间和渲染路径。",
        "JNI 批量传输、生命周期、异常和线程模型。",
        "LRU 缓存、内存复用、帧渲染和性能基准。",
        "桌面单测与 Android 集成测试如何共享。",
        "WGS84 与 GCJ-02、时间和空间数据一致性。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "17.3 必备作品材料", 2)
    add_table(
        doc,
        ["材料", "内容", "面试价值"],
        [
            ("APK", "可离线演示主路径", "证明客户端完成度"),
            ("3 分钟视频", "两个纵向场景", "快速传达用户价值"),
            ("PRD", "范围、指标、治理", "证明产品方法"),
            ("架构图", "数据到客户端全链路", "证明系统思维"),
            ("模型卡", "数据、指标、限制", "证明 AI 评估能力"),
            ("性能报告", "C++前后对比", "证明客户端深度"),
            ("GitHub 历史", "Issue、PR、CI、Release", "证明工程过程"),
        ],
        [1800, 3900, 3660],
    )

    # 18 Next actions
    add_heading(doc, "18. 下一步执行清单", 1)
    add_callout(doc, "当前状态", "产品和工程基线已经建立，下一阶段不是继续扩写方案，而是完成数据可行性和第一条纵向切片。", LIGHT_BLUE, DARK_BLUE)
    next_steps = [
        "你在 GitHub 网页创建空仓库 SkyCast，不自动生成 README/.gitignore/License。",
        "你亲手完成本地 git init、首次分批提交、remote 和 push。",
        "确定首个试点城市和一个可公开展示的晚霞历史日期。",
        "我实现公开历史事件包 Schema、读取器和测试。",
        "我实现 FastAPI 固定 Provider 与 SceneAssessment API。",
        "你安装 Android Studio、SDK、NDK；我检查环境。",
        "我实现 Android 工程壳、Room 和第一张机会卡。",
        "我们用真实用户任务完成 M2 验收，并记录第一份产品复盘。",
    ]
    add_numbered_group(doc, next_steps)

    # Appendices
    add_heading(doc, "附录 A：核心数据对象", 1)
    add_code_block(doc, [
        "ObservationFrame",
        "  source, observedAt, ingestedAt, expiresAt, crs, bbox, uri, checksum",
        "",
        "OfficialAlert",
        "  alertId, source, issuedAt, expiresAt, severity, originalText, geometry",
        "",
        "SceneAssessment",
        "  sceneType, location, validWindow, score, probability, confidence,",
        "  favorableFactors, limitingFactors, modelVersion, provenance",
        "",
        "NowcastCell",
        "  cellId, observedAt, centroid, geometry, maxReflectivity, velocity,",
        "  probabilityEnvelope, modelVersion",
        "",
        "Feedback",
        "  assessmentId, outcome, reasonCodes, occurredAt, optionalEvidence",
    ])

    add_heading(doc, "附录 B：首批 API", 1)
    add_table(
        doc,
        ["方法", "路径", "用途"],
        [
            ("GET", "/v1/cities/{cityId}/opportunities", "城市机会卡"),
            ("GET", "/v1/scenes/{sceneId}/assessment", "场景评估详情"),
            ("GET", "/v1/events/{eventId}/timeline", "历史事件时间轴"),
            ("GET", "/v1/radar/frames", "雷达帧元数据"),
            ("GET", "/v1/nowcasts/{location}", "短临概率结果"),
            ("GET", "/v1/sites", "候选地点排序"),
            ("POST", "/v1/feedback", "结果反馈"),
        ],
        [1200, 4500, 3660],
        center_cols={0},
    )

    add_heading(doc, "附录 C：参考资料与术语", 1)
    add_heading(doc, "C.1 主要参考资料", 2)
    references = [
        "[R1] 中国气象局、国家互联网信息办公室：《人工智能气象应用服务办法》（2025 年 6 月 1 日起施行）。https://www.cac.gov.cn/2025-04/30/c_1747718893685033.htm",
        "[R2] 中国气象数据网：天气雷达组网基本反射率图像产品，时间分辨率 6 分钟。https://k.data.cma.cn/mekb/?dataCode=J.0017.0010.S001&r=data/detail",
        "[R3] 国家卫星气象中心：风云卫星遥感数据服务与 FY-4 AGRI 产品。https://satellite.nsmc.org.cn/DataPortal/cn/data/dataset.html",
        "[R4] pySTEPS: an open-source Python library for probabilistic precipitation nowcasting. https://gmd.copernicus.org/articles/12/4185/2019/",
        "[R5] Android Developers：Android NDK 与 CMake。https://developer.android.com/ndk/guides/",
        "[R6] Android Developers：后台定位访问要求。https://developer.android.com/develop/sensors-and-location/location/background",
        "[R7] 高德开放平台：Android 导航 SDK 路线规划。https://lbs.amap.com/api/navigation-sdk-for-android/guide/route-plan/drive-route-plan",
        "[R8] Copernicus Data Space：Sentinel-2 数据与 API。https://dataspace.copernicus.eu/data-collections/copernicus-sentinel-missions/sentinel-2",
    ]
    for ref in references:
        add_body(doc, ref, color=MUTED)
    add_heading(doc, "C.2 术语", 2)
    add_table(
        doc,
        ["术语", "含义"],
        [
            ("Nowcasting", "通常指未来数分钟到数小时的临近预报。"),
            ("FSS", "Fractions Skill Score，用于评估空间降水/回波预测。"),
            ("Brier Score", "概率预测的均方误差，越低越好。"),
            ("Persistence", "假设当前状态保持不变的简单基线。"),
            ("Provenance", "结果所依赖的数据来源和处理血缘。"),
            ("CRS", "坐标参考系统。"),
            ("LLM Eval", "针对语言模型事实、安全、稳定和成本的系统评估。"),
        ],
        [2000, 7360],
    )

    add_callout(doc, "文档维护规则", "每次里程碑结束后更新范围、指标、风险、模型版本和下一步。产品决策改变时新增 ADR/决策记录，不直接覆盖历史原因。", LIGHT_BLUE, DARK_BLUE)

    # Core properties and save.
    doc.core_properties.title = "SkyCast AI 产品全链路规划书"
    doc.core_properties.subject = "AI 气象与天文景观决策助手产品规划"
    doc.core_properties.author = "SkyCast Project"
    doc.core_properties.keywords = "AI产品经理, Android, C++, 气象, 遥感, 雷达, 天文摄影"
    doc.core_properties.comments = "Version 1.0"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
